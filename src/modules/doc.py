from docx import Document
import os


def open_docx():
    docx_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), os.path.pardir)) + '\\static\\docx'
    docx_path = os.path.join(docx_dir, 'doc1.docx')
    document = Document(docx_path)
    return document


def handle_table(docx):
    table = docx.tables[0]
    time = table.cell(0, 2).paragraphs[0].runs[0]
    time.text = time.text.replace(time.text, '2022年3月13日15时59分')
    docx.save('../static/tmp/demo.docx')


if __name__ == '__main__':
    doc = open_docx()
    handle_table(doc)
